from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from win32com.client import constants, gencache
from docx.oxml.ns import qn


def create_doc(settings):
    """创建word文档 并设置全局样式"""
    document = Document()
    document.styles['Normal'].font.name = settings.table_font_name
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), settings.table_font_name)
    return document


def create_table(document, row_nums, col_nums, style):
    """创建表格"""
    return document.add_table(row_nums, col_nums, style)


def create_paragraph_title(document, title_text, settings):
    """创建段落标题"""
    run = document.add_heading('', level=settings.table_title_font_level).add_run(title_text)
    font_style(run, settings.table_font_name, None, settings.table_title_font_size, settings.table_content_font_color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), settings.table_font_name)


def create_paragraph(document, text, settings):
    """创建表格段落说明"""
    graph_run = document.add_paragraph().add_run(text)
    font_style(graph_run, settings.table_font_name, None, settings.table_header_font_size, settings.table_content_font_color)


def fill_table_content_cells(table, data, settings, row_nums, col_nums):
    """填充表格内容"""
    for row_num in range(1, row_nums + 1):
        for col_num in range(0, col_nums):
            table.cell(row_num, col_num).width = Inches(settings.table_col_width_dic[col_num])
            run = table.cell(row_num, col_num).paragraphs[0].add_run(str(data[row_num - 1][col_num]))
            font_style(run, settings.table_font_name,
                       settings.table_content_font_bold,
                       settings.table_content_font_size,
                       settings.table_content_font_color)


def fill_table_header_cells(table, settings):
    """填充表格标题"""
    col_nums = len(settings.table_header)
    for col_num in range(0, col_nums):
        table.cell(0, col_num).width = Inches(settings.table_col_width_dic[col_num])
        run = table.cell(0, col_num).paragraphs[0].add_run(str(settings.table_header[col_num]))
        font_style(run, settings.table_font_name,
                   settings.table_header_font_bold,
                   settings.table_header_font_size,
                   settings.table_header_font_color)


def save_doc(document, path):
    """保存文档"""
    document.save(path)


def font_style(run, font_name, bold, font_size, rgb_color):
    """设置样式"""
    if font_name:
        run.font.name = font_name
    if bold:
        run.font.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    if rgb_color:
        run.font.color.rgb = rgb_color


def doc_to_pdf(wordPath, pdfPath):
    """
    word转pdf
    :param wordPath: word文件路径
    :param pdfPath:  生成pdf文件路径
    """
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordPath, ReadOnly=1)
    doc.ExportAsFixedFormat(pdfPath,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)
